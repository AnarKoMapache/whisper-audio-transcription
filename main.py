import torch
from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline
from pydub import AudioSegment, silence
from docx import Document
import librosa
import gc
import os
import time

# Configuración del dispositivo
device = "cuda:0" if torch.cuda.is_available() else "cpu"
torch_dtype = torch.float16 if torch.cuda.is_available() else torch.float32

# Cargar modelo y procesador
model_id = "openai/whisper-medium"
model = AutoModelForSpeechSeq2Seq.from_pretrained(
    model_id, torch_dtype=torch_dtype, low_cpu_mem_usage=True
)
model.to(device)

processor = AutoProcessor.from_pretrained(model_id)
print(torch.cuda.is_available())
print(torch.cuda.device_count())
print(torch.cuda.get_device_name(0))
print(f"Modelo cargado en: {next(model.parameters()).device}")

pipe = pipeline(
    "automatic-speech-recognition",
    model=model,
    tokenizer=processor.tokenizer,
    feature_extractor=processor.feature_extractor,
    torch_dtype=torch_dtype,
    device=device,
)

# Función para convertir MP4 a WAV utilizando FFmpeg
def convert_mp4_to_wav(input_file, output_file="converted_audio.wav"):
    print("Convirtiendo MP4 a WAV...")
    try:
        audio = AudioSegment.from_file(input_file, format="mp4")
        audio = audio.set_frame_rate(16000).set_channels(1)
        audio.export(output_file, format="wav")
        print(f"Archivo convertido: {output_file}")
        return output_file
    except Exception as e:
        print(f"Error al convertir MP4 a WAV: {e}")
        return None

# Función para dividir audio dinámicamente según silencios
def split_audio_by_silence(file_path, silence_thresh=-40, min_silence_len=1000, keep_silence=300):
    print("Analizando silencios para dividir el audio...")
    try:
        audio = AudioSegment.from_file(file_path)
        chunks = silence.split_on_silence(
            audio,
            min_silence_len=min_silence_len,
            silence_thresh=silence_thresh,
            keep_silence=keep_silence
        )
        print(f"Audio dividido en {len(chunks)} fragmentos.")
        return chunks
    except Exception as e:
        print(f"Error al dividir el audio: {e}")
        return []

# Función para dividir fragmentos largos
def split_long_chunks(chunks, max_duration=30000):  # Máximo 30 segundos
    small_chunks = []
    for chunk in chunks:
        if len(chunk) > max_duration:
            small_chunks.extend(chunk[i:i + max_duration] for i in range(0, len(chunk), max_duration))
        else:
            small_chunks.append(chunk)
    print(f"Divididos fragmentos largos en un total de {len(small_chunks)} subfragmentos.")
    return small_chunks

# Función para guardar cada fragmento como archivo WAV
def save_audio_chunks(chunks, output_dir="audio_chunks"):
    os.makedirs(output_dir, exist_ok=True)
    chunk_files = []
    for i, chunk in enumerate(chunks):
        chunk_path = os.path.join(output_dir, f"chunk_{i + 1}.wav")
        chunk.export(chunk_path, format="wav")
        chunk_files.append(chunk_path)
    print(f"Fragmentos guardados en: {output_dir}")
    return chunk_files

# Función para guardar un fragmento directamente en Word
def save_chunk_to_word(doc, text):
    try:
        if doc.paragraphs:
            doc.paragraphs[-1].add_run(f" {text}")
        else:
            doc.add_paragraph(text)
        print("Fragmento transcrito y guardado.")
    except Exception as e:
        print(f"Error al guardar el fragmento: {e}")

# Función para transcribir cada archivo WAV y escribir en Word
def transcribe_audio_to_word(chunk_files, output_file="transcripcion_clase.docx"):
    try:
        doc = Document()
        doc.add_heading("Transcripción de audio", level=1)

        for i, chunk_file in enumerate(chunk_files):
            try:
                audio_array, sample_rate = librosa.load(chunk_file, sr=None)
                result = pipe({"array": audio_array, "sampling_rate": sample_rate})
                
                # Mostrar en consola cuando se procese cada fragmento
                print(f"Fragmento N° {i + 1} procesado")

                save_chunk_to_word(doc, result["text"])
            except Exception as e:
                print(f"Error al transcribir el archivo {chunk_file}: {e}")
                save_chunk_to_word(doc, "[Error en la transcripción]")
            finally:
                gc.collect()

        doc.save(output_file)
        print(f"Transcripción completa guardada en: {output_file}")
    except Exception as e:
        print(f"Error al crear el archivo Word: {e}")

# Ejecución principal
def main(mp4_audio_file):
    start_time = time.time()

    # 1. Convertir MP4 a WAV
    wav_audio_file = convert_mp4_to_wav(mp4_audio_file)
    if not wav_audio_file:
        print("Error al convertir el archivo MP4. Saliendo...")
        return

    # 2. Dividir audio en fragmentos según silencios
    audio_chunks = split_audio_by_silence(wav_audio_file)
    if not audio_chunks:
        print("No se pudo dividir el audio en fragmentos. Saliendo...")
        return

    # 3. Dividir fragmentos largos en subfragmentos
    all_chunks = split_long_chunks(audio_chunks)

    # 4. Guardar todos los fragmentos como archivos WAV
    chunk_files = save_audio_chunks(all_chunks)

    # 5. Transcribir cada archivo y guardar directamente en Word
    transcribe_audio_to_word(chunk_files)

    # 6. Medir tiempo total
    end_time = time.time()
    print(f"Tiempo total de procesamiento: {end_time - start_time:.2f} segundos")

    # Liberar recursos
    gc.collect()
    print("Recursos liberados.")

# Archivo de entrada
mp4_audio_file = "ah5.mp4"

# Ejecutar el programa
if os.path.exists(mp4_audio_file):
    main(mp4_audio_file)
else:
    print(f"El archivo {mp4_audio_file} no existe. Por favor, verifica la ruta.")
